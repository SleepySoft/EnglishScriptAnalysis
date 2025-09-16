import os
import pythoncom
import win32com.client
from pathlib import Path
from docx import Document


def remove_toc(doc):
    """
    去除Word文档中的目录
    """
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        # 判断段落是否为目录项的启发式规则
        if "TOC" in paragraph.style.name:  # 样式名包含TOC
            paragraphs_to_remove.append(paragraph)
        elif paragraph.text.strip() and (
                paragraph.text[0].isdigit() or paragraph.text.startswith(("1.", "2.", "3."))):  # 以数字或序号开头
            paragraphs_to_remove.append(paragraph)

    # 删除识别出的目录段落
    for paragraph in paragraphs_to_remove:
        p = paragraph._element
        p.getparent().remove(p)


def remove_headers_footers(doc):
    """
    删除Word文档中的页眉和页脚内容
    """
    # 遍历所有节(section)
    for section in doc.sections:
        # 清空页眉
        header = section.header
        for paragraph in header.paragraphs:
            paragraph.text = ""  # 清空段落文本
        # 清空页脚
        footer = section.footer
        for paragraph in footer.paragraphs:
            paragraph.text = ""  # 清空段落文本


def process_docx_file(file_path):
    """
    处理单个docx文件：提取正文文本和表格文本，去除目录和角色信息
    """
    doc = Document(file_path)

    # 1. 提取所有段落文本
    paragraphs_text = '\n'.join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

    # 2. 提取所有表格中的文本
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:  # 避免空单元格
                    row_text.append(cell_text)
            if row_text:  # 避免空行
                tables_text.append(' | '.join(row_text))  # 用 | 分隔同一行不同单元格内容
    tables_text_combined = '\n'.join(tables_text)

    # 3. 结合段落和表格文本
    full_text = paragraphs_text + '\n' + tables_text_combined

    return full_text


def process_all_docx_files(directory_path):
    """
    批量处理指定目录下的所有docx文件
    """
    processed_contents = {}
    for filename in os.listdir(directory_path):
        if filename.endswith('.docx'):
            file_path = os.path.join(directory_path, filename)
            try:
                content = process_docx_file(file_path)
                processed_contents[filename] = content
                print(f"成功处理: {filename}")
            except Exception as e:
                print(f"处理文件 {filename} 时出错: {str(e)}")
    return processed_contents


def convert_doc_to_docx(input_path, output_path=None):
    """
    使用 Word COM 接口将单个 .doc 文件转换为 .docx 文件

    :param input_path: 输入的 .doc 文件路径
    :param output_path: 输出的 .docx 文件路径（可选，默认为同一目录）
    :return: 成功返回 True，失败返回 False
    """
    # 初始化 COM 环境（单线程）
    pythoncom.CoInitialize()

    # 确保输入文件存在
    if not os.path.isfile(input_path):
        print(f"错误：文件不存在 - {input_path}")
        return False

    # 处理输出路径
    if output_path is None:
        output_path = str(Path(input_path).with_suffix('.docx'))
    else:
        output_path = os.path.join(output_path, os.path.basename(input_path).replace('.doc', '.docx'))

    input_path_abs = os.path.abspath(input_path)
    output_path_abs = os.path.abspath(output_path)

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    word = None
    doc = None
    try:
        # 启动 Word 应用程序
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # 后台运行，不显示界面
        word.DisplayAlerts = False  # 关闭警告提示

        # 打开 .doc 文档
        doc = word.Documents.Open(input_path_abs)

        # 另存为 .docx 格式 (FileFormat=16 代表 wdFormatXMLDocument)
        doc.SaveAs2(output_path_abs, FileFormat=16)

        print(f"转换成功: '{os.path.basename(input_path_abs)}' -> {os.path.basename(output_path_abs)}")
        return True

    except Exception as e:
        print(f"转换失败: '{os.path.basename(input_path_abs)}' - 错误: {str(e)}")
        return False

    finally:
        # 确保无论如何都尝试释放资源
        try:
            if 'doc' in locals():
                doc.Close(SaveChanges=False)
        except:
            pass
        try:
            if 'word' in locals():
                word.Quit()
        except:
            pass


def batch_convert_doc_to_docx(input_dir, output_dir=None, recursive=True):
    """
    批量转换目录中的所有 .doc 文件

    :param input_dir: 包含 .doc 文件的输入目录
    :param output_dir: 输出目录（可选，默认为输入目录）
    :param recursive: 是否递归处理子目录
    """
    if output_dir is None:
        output_dir = input_dir

    # 收集所有 .doc 文件
    if recursive:
        doc_files = list(Path(input_dir).rglob("*.doc"))
    else:
        doc_files = list(Path(input_dir).glob("*.doc"))

    total_files = len(doc_files)
    if total_files == 0:
        print("在指定目录中未找到 .doc 文件")
        return

    print(f"找到 {total_files} 个 .doc 文件，开始转换...")

    success_count = 0
    # 遍历并转换每个文件
    for i, doc_path in enumerate(doc_files, 1):
        print(f"正在处理 ({i}/{total_files}): {doc_path.name}")
        if convert_doc_to_docx(str(doc_path), output_dir):
            success_count += 1

    # 输出转换报告
    print("\n" + "=" * 50)
    print(f"批量转换完成！")
    print(f"成功: {success_count} 个")
    print(f"失败: {total_files - success_count} 个")
    print("=" * 50)
