import nltk
from nltk import pos_tag
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from collections import Counter
import string
import pandas as pd


# pip install python-docx nltk matplotlib pandas openpyxl

# 确保已下载所需的NLTK数据
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('punkt_tab')
nltk.download('wordnet')
nltk.download('averaged_perceptron_tagger')  # 用于词性标注，帮助词形还原
nltk.download('averaged_perceptron_tagger_eng')


def analyze_collocations(text, top_n=20):
    """
    分析常见的词性搭配模式，这有助于发现英语中的习惯用法[6](@ref)。
    例如：动词+介词（VB+IN）、形容词+名词（JJ+NN）等。
    """

    tokens = word_tokenize(text)
    tagged_tokens = pos_tag(tokens)

    # 定义一些常见的、有意义的词性组合模式

    patterns = [
        # 基础模式
        (('VB', 'IN'), 'Verb+Prep'),  # 动词+介词，如：look at, depend on, talk about
        (('VB', 'DT', 'NN'), 'Verb+Det+Noun'),  # 动词+限定词+名词，如：have a look, make a decision, take the chance
        (('JJ', 'NN'), 'Adj+Noun'),  # 形容词+名词，如：red apple, important meeting, difficult situation
        (('RB', 'VB'), 'Adv+Verb'),  # 副词+动词，如：quickly run, easily understand, carefully consider
        (('NN', 'IN', 'NN'), 'Noun+Prep+Noun'),  # 名词+介词+名词，如：transition to adulthood, key to success, fear of failure
        (('VB', 'RB'), 'Verb+Adv'),  # 动词+副词，如：speak clearly, work efficiently, respond immediately
        (('IN', 'DT', 'NN'), 'Prep+Det+Noun'),  # 介词+限定词+名词，如：in the morning, on a mission, with an idea
        (('NN', 'NN'), 'Compound Noun'),  # 复合名词，如: coffee cup, business meeting, research paper
        (('VB', 'NN'), 'Verb+Noun'),  # 动词+名词，如: make progress, take notes, set goals

        # 新增模式
        (('VB', 'DT', 'JJ', 'NN'), 'Verb+Det+Adj+Noun'),
        # 动词+限定词+形容词+名词，如：have a great day, make an important decision, see the beautiful sunset
        (('JJ', 'JJ', 'NN'), 'Adj+Adj+Noun'),  # 形容词+形容词+名词，如：beautiful red rose, large wooden table, small black cat
        (('RB', 'JJ'), 'Adv+Adj'),  # 副词+形容词，如：extremely important, very happy, quite difficult
        (('NN', 'VB'), 'Noun+Verb'),  # 名词+动词，如：problem solving, decision making, time management
        (('VB', 'PRP'), 'Verb+Pronoun'),  # 动词+代词，如：help me, tell them, ask us
        (('IN', 'JJ', 'NN'), 'Prep+Adj+Noun'),  # 介词+形容词+名词，如：in great detail, with special care, on important matters
        (('DT', 'NN', 'IN', 'NN'), 'Det+Noun+Prep+Noun'),
        # 限定词+名词+介词+名词，如：the end of time, a piece of cake, the beginning of history
        (('MD', 'VB', 'RB'), 'Modal+Verb+Adv'),
        # 情态动词+动词+副词，如：can easily do, will quickly go, should carefully consider
        (('NN', 'IN', 'DT', 'NN'), 'Noun+Prep+Det+Noun'),
        # 名词+介词+限定词+名词，如：transition to a new, solution to the problem, key to a mystery
        (('VB', 'TO', 'VB'), 'Verb+To+Verb'),  # 动词+不定式标记+动词，如：want to go, need to see, try to understand
        (('VBG', 'NN'), 'Gerund+Noun'),  # 动名词+名词，如：reading books, making progress, writing letters
        (('VBN', 'IN'), 'PastPart+Prep'),  # 过去分词+介词，如：interested in, covered with, known for
        (('CD', 'NNS'), 'Number+PluralNoun'),  # 基数词+名词复数，如：three books, five years, ten students
        (('JJ', 'CC', 'JJ'), 'Adj+Conj+Adj'),  # 形容词+连词+形容词，如：simple and effective, short but clear, tired yet happy
        (('VB', 'PRP', 'RB'), 'Verb+Pronoun+Adv'),  # 动词+代词+副词，如：tell me quickly, show them clearly, ask us politely
        (('RB', 'RB', 'JJ'), 'Adv+Adv+Adj'),
        # 副词+副词+形容词，如：very extremely hot, quite surprisingly good, rather unexpectedly cold
        (('DT', 'JJ', 'NN', 'VBZ'), 'Det+Adj+Noun+Verb'),
        # 限定词+形容词+名词+动词，如：the quick brown fox jumps, a beautiful red rose blooms
        (('PRP', 'MD', 'VB', 'RB'), 'Pron+Modal+Verb+Adv'),
        # 代词+情态动词+动词+副词，如：I can easily do, you should carefully consider, we will quickly go
        (('NN', 'VBZ', 'JJ'), 'Noun+Verb+Adj'),  # 名词+动词+形容词，如: time flies fast, sun sets red, water runs clear
        (('IN', 'PRP$', 'NN'), 'Prep+Possessive+Noun')  # 介词+物主代词+名词，如：in my opinion, on his behalf, with her permission
    ]

    collocation_counts = {desc: Counter() for _, desc in patterns}
    window_size = 4  # 检查相邻单词的窗口大小

    for i in range(len(tagged_tokens) - window_size + 1):
        window = tagged_tokens[i:i + window_size]
        for (pattern, description) in patterns:
            if len(pattern) <= len(window):
                match = True
                for j in range(len(pattern)):
                    if window[j][1] != pattern[j]:
                        match = False
                        break
                if match:
                    phrase = ' '.join([word for word, pos in window[:len(pattern)]])
                    collocation_counts[description][phrase] += 1

    # 获取每种模式的前top_n个最常见搭配
    top_collocations = {}
    for desc, counter in collocation_counts.items():
        top_collocations[desc] = counter.most_common(top_n)

    return top_collocations


def preprocess_and_analyze(text):
    # 1. 分句
    sentences = sent_tokenize(text)

    # 2. 分词 & 预处理（转换为小写、去除标点和停用词）
    stop_words = set(stopwords.words('english'))
    translator = str.maketrans('', '', string.punctuation)
    all_words = []

    for sentence in sentences:
        words = word_tokenize(sentence)
        # 转换为小写并去除标点符号
        words = [word.translate(translator).lower() for word in words]
        # 去除停用词和空字符串
        words = [word for word in words if word not in stop_words and word]
        all_words.extend(words)

    # 3. 词形还原 (需要先进行词性标注以获得最佳效果)
    # 注意：为简化示例，我们假设所有词都是名词('n')。实际应用中应进行词性标注。
    lemmatizer = WordNetLemmatizer()
    lemmatized_words = [lemmatizer.lemmatize(word, pos='v') for word in all_words]  # 尝试动词形式

    # 4. 统计词频
    word_freq = Counter(lemmatized_words)
    return sentences, word_freq


# # 示例用法
# your_text = "Your sample text goes here. It can contain multiple sentences. Words like running, ran, and runs should be grouped."
# sentences, frequency = preprocess_and_analyze(your_text)
#
# print("词频统计结果 (合并变形后):")
# for word, count in frequency.most_common(10):  # 打印最常见的10个词
#     print(f"{word}: {count}")


def analyze_sentence_patterns(sentences):
    """
    一个简单基于词性标签序列的句型统计示例。
    这只是一个初级方法，更准确的句型分析需要依赖句法分析。
    """
    # 常见的词性标记: NN(名词), VB(动词), IN(介词), DT(冠词), JJ(形容词), PRP(人称代词)
    pattern_freq = Counter()

    for sentence in sentences:
        words = word_tokenize(sentence)
        # 获取每个词的词性标签
        pos_tags = [tag for word, tag in nltk.pos_tag(words)]
        # 将词性标签序列转换为一个字符串，作为句型的近似表示
        pattern_key = ' '.join(pos_tags)
        pattern_freq[pattern_key] += 1

        # 你也可以定义一些规则，将特定的词性序列映射到你定义的句型名称上
        # if pos_tags starts with 'PRP VB' -> S+V pattern?

    return pattern_freq


# # 示例用法
# sentence_pattern_freq = analyze_sentence_patterns(sentences)
#
# print("\n句型（词性序列）统计结果:")
# for pattern, count in sentence_pattern_freq.most_common():
#     print(f"{pattern}: {count}")

import re
from docx import Document
import os


def remove_toc(doc):
    """
    去除Word文档中的目录
    """
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        # 判断段落是否为目录项的启发式规则
        if "TOC" in paragraph.style.name:  # 样式名包含TOC
            paragraphs_to_remove.append(paragraph)
        elif paragraph.text.strip() and (
                paragraph.text[0].isdigit() or paragraph.text.startswith(("1.", "2.", "3."))):  # 以数字或序号开头
            paragraphs_to_remove.append(paragraph)

    # 删除识别出的目录段落
    for paragraph in paragraphs_to_remove:
        p = paragraph._element
        p.getparent().remove(p)


def remove_headers_footers(doc):
    """
    删除Word文档中的页眉和页脚内容
    """
    # 遍历所有节(section)
    for section in doc.sections:
        # 清空页眉
        header = section.header
        for paragraph in header.paragraphs:
            paragraph.text = ""  # 清空段落文本
        # 清空页脚
        footer = section.footer
        for paragraph in footer.paragraphs:
            paragraph.text = ""  # 清空段落文本


def remove_role_info(text):
    """
    去除正文中的角色信息（例如：Peppa: xxx）
    使用正则表达式匹配并移除角色名称及冒号
    """
    # 匹配模式：角色名（可能包含空格和特殊字符）后跟冒号和可选空格
    pattern = r'^[A-Za-z\s]+:\s*'
    # 逐行处理，移除匹配的角色信息
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        cleaned_line = re.sub(pattern, '', line)
        cleaned_lines.append(cleaned_line)
    return '\n'.join(cleaned_lines)


def remove_specific_patterns(text):
    """
    从文本中删除特定模式的字符串。
    该模式通常包含：中文文字 + 制表符（\\t） + 分数数字（如36/39） + 可能的换行符（\\n）

    参数:
    text (str): 需要处理的原始文本

    返回:
    str: 处理后的纯净文本
    """
    # 定义正则表达式模式，匹配中文文字、制表符、分数数字和可能的换行符
    # 模式解释：
    # [\u4e00-\u9fa5]+ : 匹配一个或多个中文字符
    # \t+             : 匹配一个或多个制表符
    # \d+/\d+         : 匹配分数形式的数字（如36/39）
    # \n?             : 匹配0个或1个换行符（可能出现在模式末尾）
    pattern = r'[\u4e00-\u9fa5]+\t+\d+/\d+\n?'

    # 使用re.sub将匹配到的模式替换为空字符串
    cleaned_text = re.sub(pattern, '', text)

    return cleaned_text


def process_docx_file(file_path):
    """
    处理单个docx文件：提取正文，去除目录和角色信息
    """
    # 加载Word文档
    doc = Document(file_path)

    # 1. 去除目录
    remove_toc(doc)
    remove_headers_footers(doc)

    # 2. 提取正文文本（连接所有段落）
    full_text = '\n'.join([para.text for para in doc.paragraphs])

    # 3. 去除角色信息
    cleaned_text = remove_role_info(full_text)
    cleaned_text = remove_specific_patterns(cleaned_text)

    return cleaned_text


def process_all_docx_files(directory_path):
    """
    批量处理指定目录下的所有docx文件
    """
    processed_contents = {}
    for filename in os.listdir(directory_path):
        if filename.endswith('.docx'):
            file_path = os.path.join(directory_path, filename)
            try:
                content = process_docx_file(file_path)
                processed_contents[filename] = content
                print(f"成功处理: {filename}")
            except Exception as e:
                print(f"处理文件 {filename} 时出错: {str(e)}")
    return processed_contents


# 使用示例
if __name__ == "__main__":
    # directory = "Docs"
    # results = process_all_docx_files(directory)
    #
    # # 查看处理结果（例如，打印第一个文件的内容）
    # with open('pure_text.txt', 'wt') as f:
    #     for filename, content in results.items():
    #         print(f"文件: {filename}")
    #         print("处理后的内容预览:")
    #         print(content[:500] + "..." if len(content) > 500 else content)  # 打印前500字符
    #         print("\n" + "=" * 50 + "\n")
    #         f.write(content)

    with open('pure_text.txt', 'rt') as f:
        full_text = f.read()

    # ------------------------------------------------------------------------------------------------------------------

    sentences, frequency = preprocess_and_analyze(full_text)

    print("词频统计结果 (合并变形后):")
    for word, count in frequency.most_common(10):  # 打印最常见的10个词
        print(f"{word}: {count}")

    # 创建句子列表的DataFrame
    df_sentences = pd.DataFrame(sentences, columns=['Sentences'])

    # 创建词频统计的DataFrame
    df_frequency = pd.DataFrame(frequency.items(), columns=['Word', 'Frequency'])
    # 按词频从高到低排序
    df_frequency.sort_values(by='Frequency', ascending=False, inplace=True)

    with pd.ExcelWriter('text_analysis_results.xlsx', engine='openpyxl') as writer:
        df_sentences.to_excel(writer, sheet_name='Sentences', index=False)
        df_frequency.to_excel(writer, sheet_name='Word Frequency', index=False)

    print("分析结果已成功导出到 'text_analysis_results.xlsx'")

    # ------------------------------------------------------------------------------------------------------------------

    collocations = analyze_collocations(full_text)

    print("\n=== 常见表达方式归类 ===")
    for pattern_type, phrases in collocations.items():
        if phrases:  # 只显示有结果的类型
            print(f"\n--- {pattern_type} ---")
            for phrase, freq in phrases:
                print(f"{phrase}: {freq}")

    data = []
    for collocation_type, phrases_list in collocations.items():
        for phrase, frequency in phrases_list:
            data.append({
                "搭配模式": collocation_type,
                "搭配短语": phrase,
                "出现频率": frequency
            })

    df_collocations = pd.DataFrame(data)

    # 3. 写入 Excel 文件
    excel_filename = "collocation_analysis_results.xlsx"
    df_collocations.to_excel(excel_filename, index=False)

    print(f"搭配分析结果已保存到 '{excel_filename}'")
